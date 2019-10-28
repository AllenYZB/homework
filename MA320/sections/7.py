# -*- encode: utf-8 -*-
from docx import Document
from docx.shared import Inches


document = Document()
indent = '    '

# section
document.add_heading('Summary', 0)

document.add_paragraph(indent + 'This chapter describes the mechanics of the publication process, which includes choosing a journal, submitting a manuscript, and the refereeing process. Strictly, your article can be called "manuscript" only after submitting, and "paper" only after it has been accepted for publication. Moreover, the textbook proposes a valuable point: do not post articles that are not worth publishing.')

document.add_heading('Choosing a Journal', 1)
document.add_paragraph(indent + 'Submission is the next step after the article has been written and repeatedly revised, and choosing an appropriate journal is the most important thing before submission, which requires to consider the following factors: the length of the review time and the interval from receipt to publication. Furthermore, the textbook tell us that do not pursue the reputation of a journal: the value of a paper is not being recognized when it is published, but being continually cited later.')

document.add_heading('Submitting a Manuscript', 1)
p = document.add_paragraph(indent + 'When finishing choosing the journal, you will enter the process of submission. Before submitting your manuscript, you should read carefully the instructions for authors that are printed in each issue of your chosen journal. Then you should write a cover letter and deliver it with your manuscript to the chosen journal. In addition, most of the requirements are common sense and are similar for each journal, please see page 129 in ')
p.add_run('Handbook of Writing for the Mathematical Science').italic = True
p.add_run(' for more information.')

document.add_heading('The Refereeing Process', 1)
document.add_paragraph(indent + 'If a manuscript is submitted to "The Editor" of the chosen journal, it is passed on to the editor-in-chief, who assigns it to a member of the editorial board. Then the chosen editor writes to two or more people asking them to referee the manuscript, suggesting a deadline about six weeks from the time they receive the manuscript. Finally, when all the referee reports have been received, the editor decides the fate of the manuscript, informs the author, and notifies the journal.')
document.add_paragraph(indent + 'If you receive a rejection letter for any reason, you should treat it objectively. At this point, you have three options: 1) When the content of your article does not meet the chosen journal field, consider submitting it to other journals; 2) Modify your article according to the review comments and then submit it again; 3) Write a rebuttal to the editor to prove why the reviewer is wrong.')



# generate docx file
document.save('7.docx')
