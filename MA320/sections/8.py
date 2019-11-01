# -*- encode: utf-8 -*-
from docx import Document
from docx.shared import Inches


document = Document()
indent = '    '

# section
document.add_heading('Summary of Two Textbooks', 0)
document.add_heading('The Purpose of a Thesis', 1)
document.add_paragraph(indent + 'The purpose of a thesis varies with the type of degree and the institution. For the master\'s degree, the thesis has to satisfy the following criteria:')
document.add_paragraph('show that the student has read and understood a body of research literature;', style='List Bullet')
document.add_paragraph('provide evidence that the student is capable of carrying out original research.', style='List Bullet')
document.add_paragraph(indent + 'For the Ph.D., the thesis has to additionally satisfy the following criteria:')
document.add_paragraph('represent a significant contribution to the field.', style='List Bullet')
document.add_paragraph(indent + 'As for the institution, it is worth checking what is expected by your institution.')

document.add_heading('Content', 1)
document.add_paragraph(indent + 'First of all, a thesis must be "self-contained", which means that is must stand on its own as a complete account of the author\'s work on the subject of investigation; Secondly, a thesis is formatted like a book with probably more than one topic, which makes it usually longer than an average paper; Last but not least, the primary readers of a thesis will read it with at least as much care as do the referees of a paper, which means that the primary readers are its judges.')
document.add_paragraph(indent + 'In a thesis, you should generally include details, instead of padding with unnecessary material. And a thesis must have a fairly rigid structure, for example, the problem being addressed must be clearly described and put into context in the first one or two chapters; conclusions must be carefully draw and the overall contribution of the thesis assessed at the end of the thesis. Moreover, it is important to avoid inadvertently committing plagiarism (do not learn this article, since it is a summary.).')
document.add_paragraph(indent + 'The advice of when to write the thesis is to start sooner rather than later. In the first few months of study, you will become familiar with the problem and the literature, then you can begin to draft the first few chapters and collect references for the bibliography.')

document.add_heading('Presentation', 1)
document.add_paragraph(indent + 'The rules about the presentation of a thesis vary with the institution, therefore, it is worth checking what form or binding are regulated by your institution. Furthermore, you can inspect to check the required format of the previous successful thesis contained in library.')



# generate docx file
document.save('8.docx')
