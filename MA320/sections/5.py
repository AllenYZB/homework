# -*- encode: utf-8 -*-
from docx import Document
from docx.shared import Inches


document = Document()
indent = '    '

# section
document.add_heading('Summary', 0)
# for section in document.sections:
# 	for paragraph in section.header.paragraphs:
# 		paragraph.text = '11711217, Iydon Liang'



# generate docx file
document.save('5.docx')
