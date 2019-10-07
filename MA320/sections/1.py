# -*- encode: utf-8 -*-
from docx import Document
from docx.shared import Inches


document = Document()
indent = '    '

# section
document.add_heading('Summary', 0)
# for section in document.sections:
# 	for paragraph in section.header.paragraphs:
# 		paragraph.text = '11711217, Iydon Liang'

document.add_paragraph('Mathematical papers have several types, such as research paper, survey article (review paper), referee report, research proposal, etc. Each type of paper has its characteristics, but they have commonalities in terms of writing. Before writing the paper, we need to think about its high-level architecture carefully, which is a good idea to reduce workload and improve work efficiency. When talking about architecture, there are several aspects that we need to pay attention to the target audience, related materials, paper organization, and repeated modification.')


# subsection
document.add_heading('Target Audience', 1)
document.add_paragraph('The first task in writing a paper is to target the audience, which helps to determine the writing style, content layout, and article focus by considering the breadth of the intended readership. Identifying the audience can be useful when finishing the paper, you can check whether the paper is well-focused by asking yourself why your intended audience should want to read the paper. If this question can be easily answered, it means that your paper concentrates on the main theme; else you might consider altering the aims or re-working on the current result to unify the theme of the paper.')


# subsection
document.add_paragraph('Collecting some writing-related materials, especially with a few closely related documents, can help you not only in understanding the history and development of your research project but also in using the terminology appropriately. Therefore you can highlight keywords and key sentences in other documents, which can be referenced in future writing. Plagiarism is not allowed under any circumstances, but you can just choose a few words or sentences from references and rephrase them under your understanding.')


# subsection
document.add_heading('Paper Organization', 1)
p = document.add_paragraph('A standard mathematical paper is generally composed of ')
p.add_run('title, abstract, introduction, the main body of the article, conclusions, acknowledgments, references, and appendix').italic = True
p.add_run('. Furthermore, they all have a specific purpose.')
# subsubsection
document.add_heading('Title', 2)
document.add_paragraph('The title is to describe the main theme of the paper with the least words and the largest information to attract the audience to further read the abstract or the main body of the paper.')

# subsubsection
document.add_heading('Abstract', 2)
document.add_paragraph('The abstract is to briefly inform the reader of the main content of this paper with the main purpose, thoughts, and results of the paper to keep the interest in reading the main body.')

# subsubsection
document.add_heading('Introduction', 2)
document.add_paragraph('The introduction should provide a comprehensive, accurate and objective characterization to the background and the development of the issues, which will be discussed in the paper, to give the audience valuable information and good impressions.')

# subsubsection
document.add_heading('Main Body', 2)
document.add_paragraph('The main body is the most important part of the paper, it comprehensively demonstrates the research results and the analysis of results. The main body of the paper can be composed of various components, which leads to several structures, such as,')

document.add_paragraph('preliminaries, main results, the outline of proofs and extensions;', style='List Bullet')
document.add_paragraph('the problem or governing equations, the numerical method or experimental method, theoretical analysis, numerical results, and conclusions.', style='List Bullet')

# subsubsection
document.add_heading('Conclusions', 2)
document.add_paragraph('The conclusions should provide a summary of the main contribution, explain the questions mentioned in the introduction, point out other aspects or broader issues that the paper does not consider, talk about the possible impact of the research and identify the next step or look to the future.')

# subsubsection
document.add_heading('Acknowledgments', 2)
document.add_paragraph('Acknowledgments are mainly for the people or institutions which have provided help to paper writing. Moreover, we must thank the institutions that fund the research in most cases.')

# subsubsection
document.add_heading('References', 2)
document.add_paragraph('The papers and books which are related to the content or cited in the paper should be listed under the name of the references, thus the audience can judge whether the information is familiar and whether it is in line with her interest.')

# subsubsection
document.add_heading('Appendix', 2)
document.add_paragraph('On the one hand, the appendix satisfies the audience who wants to fully understand the process of theorem inference, on the other hand, the appendix provides some new ideas for those who do not see the proof temporarily. Moreover, the author can add the materials, which are useful for the results, to the appendix, and the reader can do the future work through the appendix. It is a win-win setting.')

# subsubsection
document.add_heading('Repeated Modification', 2)
document.add_paragraph('When completing the first draft of the paper, it is necessary to repeatedly modify it. When writing the first English paper in life, you should find some people, who have some English writing experience or are native speakers, to modify it again and again. You should carefully ponder the reasons why others modify these sentences, in order to improve your writing level.')



# generate docx file
document.save('1.docx')
