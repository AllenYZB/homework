# -*- encode: utf-8 -*-
from docx import Document
from docx.shared import Inches


document = Document()
indent = '    '

# section
document.add_heading('Summary', 0)
# for section in document.sections:
# 	for paragraph in section.header.paragraphs:
# 		paragraph.text = '11711217, Iydon Liang'

document.add_paragraph('If English is not your first language then writing the mathematical paper is doubly difficult, since you cannot guarantee the proper use of words even though with a large scientific and technical vocabulary. Therefore, Halmos comes up with three specific advice on how to use scientific and technical terms carefully,')
document.add_paragraph('Avoid using technical terms as much as possible, especially creating new ones;', style='List Bullet')
document.add_paragraph('Be careful when thinking about the new terms you have to create by looking up the dictionary to make it as appropriate as possible;', style='List Bullet')
document.add_paragraph('Correctly and consistently use the old terms without awkwardness.', style='List Bullet')

document.add_paragraph(indent + 'In addition to the terminology, you should also pay attention to the language itself. Science and technology English is generally not emotional, and its expression is relatively straightforward, which helps to make the readers easy to understand without creating too much imagination irrelevant. That is, the vocabulary using in science and technology English requires specific and sTable. Therefore, you need to accumulate basic mathematical vocabulary, mathematical symbols with their pronunciation, common phrases and language structures to better write rational mathematics papers.')


# subsection
document.add_heading('Basic Vocabulary', 1)
document.add_paragraph('Basic vocabulary is an important prerequisite because you have to grapple with a large vocabulary as you try to express your thoughts. There is a large mathematical vocabulary, but the spelling is relatively simple compared to other disciplines, and as long as you master a certain amount of mathematical vocabulary, writing an uncomplicated mathematics paper will not be too difficult.')
p = document.add_paragraph(indent + 'Please see Table 2.1 to Table 2.13 in ')
p.add_run('Mathematical Writing in English').italic = True
p.add_run('[1] about the basic vocabulary in different fields of mathematics.')


# subsection
document.add_heading('Symbol Overview', 1)
document.add_paragraph('Mathematical language is the universal language, therefore, when using English to write technical papers or using spoken English to express certain symbols in scientific communication, you should understand the English expression of the common mathematical symbols. In this way, you can do more with less.')
p = document.add_paragraph(indent + 'Please see Table 2.14 to Table 2.17 in ')
p.add_run('Mathematical Writing in English').italic = True
p.add_run('[1] about the common symbols in different fields of mathematics.')


# subsection
document.add_heading('Common Phrases', 1)
document.add_paragraph('The common mathematical phrases found in professional English textbooks or theoretically papers are very practical and require proper collection and organization because proficiently using these phrases can be beneficial to your writing.')
p = document.add_paragraph(indent + 'Please see Table 2.18 to Table 2.20 in ')
p.add_run('Mathematical Writing in English').italic = True
p.add_run('[1] about the common phrases.')


# subsection
document.add_heading('Language Structures', 1)
document.add_paragraph('Since mathematics writing has certain rules and formats in language structures and styles, there are many fixed modes in mathematics writing, which include some common sentence patterns, modified vocabulary, verbs in mathematical operations and transition statements.')
p = document.add_paragraph('Please see Table 2.21 in ', style='List Bullet')
p.add_run('Mathematical Writing in English').italic = True
p.add_run('[1] about the common sentence patterns;')
p = document.add_paragraph('Please see Table 2.22 in ', style='List Bullet')
p.add_run('Mathematical Writing in English').italic = True
p.add_run('[1] about the modified vocabulary;')

p = document.add_paragraph('Please see Table 2.23 in ', style='List Bullet')
p.add_run('Mathematical Writing in English').italic = True
p.add_run('[1] about the verbs used in mathematical operations;')

p = document.add_paragraph('Please see Table 2.24 in ', style='List Bullet')
p.add_run('Mathematical Writing in English').italic = True
p.add_run('[1] about the common verbs used in mathematics writing;')

p = document.add_paragraph('Please see Section 2.4.5 in ', style='List Bullet')
p.add_run('Mathematical Writing in English').italic = True
p.add_run('[1] about the transition statements.')


# subsection
document.add_heading('Proof Expression', 1)
document.add_paragraph('Mathematical proof is relatively easy in mathematics writing, but it requires clear and concise expression as well. Therefore, you can express mathematical proof skillfully after an amount of training.')

document.add_paragraph(indent + 'Some examples are often used in proof[2],')

document.add_paragraph('At the beginning of the proof, "you should strive to keep the reader informed of where you are in the proof and what remains to be done." Useful phrases include', style='List Bullet')
document.add_paragraph('First, we establish that ...', style='List Bullet 2')
document.add_paragraph('Our task is now to ...', style='List Bullet 2')
document.add_paragraph('Our problem reduces to ...', style='List Bullet 2')
document.add_paragraph('It remains to show that ...', style='List Bullet 2')
document.add_paragraph('We are almost ready to invoke ...', style='List Bullet 2')

document.add_paragraph('At the end of the proof, it is often marked by the halmos symbol □. Sometimes the abbreviation QED is used instead.', style='List Bullet')

document.add_paragraph('If you omit part of a proof, "it is best to indicate the nature and length of the omission", via phrases such as the following.', style='List Bullet')
document.add_paragraph('It is easy/simple/straightforward to show that ...', style='List Bullet 2')
document.add_paragraph('Some tedious manipulation yields  ...', style='List Bullet 2')
document.add_paragraph('An easy/obvious induction gives ...', style='List Bullet 2')
document.add_paragraph('After two applications of ... we find ...', style='List Bullet 2')
document.add_paragraph('An argument similar to the ones used in ... shows that ...', style='List Bullet 2')


# subsection
document.add_heading('Important Conjunction', 1)
document.add_paragraph('"Join clauses good, like a conjunction should." This sentence points out the importance of conjunctions in English writing, thus, we should be cautious when using conjunctions. This subsection explains the different functions of conjunctions through a series of examples, including the following functions,')
document.add_paragraph('Combinations:', style='List Bullet')
document.add_paragraph('and, both, also, as well as, not only ... but also, apart from, in addition to, moreover, furthermore;', style='List Bullet 2')
document.add_paragraph('Implications or explanations;', style='List Bullet')
document.add_paragraph('as, because, since, due to, in view of, owing to, on account of, given, it follows that, consequently, therefore, thus;', style='List Bullet 2')
document.add_paragraph('Modifications and restrictions;', style='List Bullet')
document.add_paragraph('alternatively, although, though, but, whereas, by contrast, except, however, on the other hand, nevertheless, despite, in spite of, instead of, rather that.', style='List Bullet 2')


# subsection
document.add_heading('Comparison', 1)
document.add_heading('Symbols In the Sentences', 2)
document.add_paragraph('Halmos once said, "the best notation is no notation". That is, symbols should be used as little as possible when writing sentences that include mathematical symbols unless you have to.')
p = document.add_paragraph(indent + 'Please see Section 2.7.1 in ')
p.add_run('Mathematical Writing in English').italic = True
p.add_run('[1] for more examples.')

document.add_heading('Definite Article and Indefinite Article', 2)
document.add_paragraph('It is difficult for speakers, who speak the languages that do not have articles, to use articles correctly in English. Therefore, there are two important rules[2],')

p = document.add_paragraph('Do not use ', style='List Bullet')
p.add_run('the').italic = True
p.add_run(' (with plural or uncountable nouns) to talk about things in general;')
document.add_paragraph('Do not use singular countable nouns without articles.', style='List Bullet')
p = document.add_paragraph(indent + 'Please see Section 2.7.2 in ')
p.add_run('Mathematical Writing in English').italic = True
p.add_run('[1] for more examples.')

document.add_heading('Common Misspellings or Confusions', 2)
document.add_paragraph('We often find that many words are spelling similar when learning English words, but we are not sure whether the words are used properly during the writing process. Thus we need software and accurate memory to reduce misspellings and confusion.')
p = document.add_paragraph(indent + 'Please see Table 2.25 to Table 2.26 in ')
p.add_run('Mathematical Writing in English').italic = True
p.add_run('[1] for more examples.')


# references
document.add_heading('References', 1)
document.add_paragraph('汤涛, 丁玖. 数学之英文写作[M]. 北京: 高等教育出版社, 2013.', style='List Number')
document.add_paragraph('HIGHAM N J. Handbook of writing for the mathematical sciences[M]. Philadelphia: Society for Industrial and Applied Mathematics, 1998.', style='List Number')



# generate docx file
document.save('3+4.docx')
